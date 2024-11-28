from docx import Document
from io import BytesIO
from typing import List, NoReturn

from documents import word_api
from utils import security

def create_word_user(document: Document, data: List[str]) -> NoReturn:
    
    """
    Creates word document with blanks
    """

    alphabet = "ABCDEFGHIJKLMNOPQRSTUVWXYZ                                             "

    names = [
        "\nФИО:",
        "Класс:",
        "Дата:",
        "Тема:\n"
    ]

    data[0] = '\n' + data[0]
    data[2] += " г."

    section = document.sections[0]
    section = word_api.set_margin(section)
    section = word_api.create_borders(section)

    header_table = word_api.create_table(document, 4, 3)
    header_table = word_api.set_column_width(header_table, 0, 3.75)
    header_table = word_api.set_column_width(header_table, 1, 8.7)
    header_table = word_api.set_column_width(header_table, 2, 7)
    word_api.merge_cells(list(header_table.columns[2].cells))

    name_cells = header_table.columns[0].cells
    value_cells = header_table.columns[1].cells
    image_cell = header_table.columns[2].cells[0]

    for i, name_cell in enumerate(name_cells):
        name_paragraph = name_cell.paragraphs[0]
        word_api.align_paragraph(name_paragraph, "right")
        word_api.write_text_to_paragraph(name_paragraph, f"{names[i]}")

    for i, value_cell in enumerate(value_cells):
        value_paragraph = value_cell.paragraphs[0]
        word_api.align_paragraph(value_paragraph, "left")
        word_api.write_text_to_paragraph(value_paragraph, 
                                         f"{data[i]}",
                                         underline = True)

    secret_code = security.create_secret_code()
    qrcode = security.create_qrcode(secret_code)

    image_paragraph = image_cell.paragraphs[0]
    word_api.align_paragraph(image_paragraph, "right")
    word_api.add_qrcode(image_paragraph, qrcode)

    paragraph = document.add_paragraph()
    word_api.write_text_to_paragraph(paragraph, 
                                     '',
                                     size = 26)
    
    answers_table = word_api.create_table(document, 14, 4)
    answers_table.style = "Table Grid"
    word_api.align_table(answers_table, "center")

    for i in range(len(answers_table.rows)):
        word_api.set_row_height(answers_table, i, 35)

    for i in range(len(answers_table.columns)):
        word_api.set_column_width(answers_table, i, 4.5)

    merged1 = word_api.merge_cells(list(answers_table.rows[-1].cells[-2::]))
    merged2 = word_api.merge_cells(list(answers_table.rows[-2].cells[-2::]))
    word_api.merge_cells([merged1, merged2])

    for i in range(len(answers_table.rows)):
        if (i % 2 == 0):
            for j, cell in enumerate(answers_table.rows[i].cells):
                letter_paragraph = cell.paragraphs[0]
                word_api.align_paragraph(letter_paragraph, "center")
                word_api.write_text_to_paragraph(letter_paragraph, 
                                                 alphabet[i // 2 * 4 + j],
                                                 bold = False,
                                                 size = 26)
        else:
            for cell in answers_table.rows[i].cells:
                space_paragraph = cell.paragraphs[0]
                word_api.align_paragraph(space_paragraph, "center")
                word_api.write_text_to_paragraph(space_paragraph,
                                                 ' ',
                                                 size = 28)

def create_word_users(data: List[List[str]]) -> BytesIO:
    document = Document()

    for user in data:
        create_word_user(document, user)

    f = BytesIO()
    document.save(f)
    f.seek(0)

    return f