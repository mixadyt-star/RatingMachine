from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.shared import Cm, Pt
from docx.section import Section
from docx.oxml import OxmlElement
from docx.oxml.shared import qn
from typing import NoReturn, List
from io import BytesIO

def create_borders(section: Section, 
                   border_type: str = "single",
                   border_size: int | str = 20,
                   border_space: int | str = 24,
                   border_color: str = "black") -> Section:
    """
    Create borders in section
    """

    section_properties = section._sectPr
    page_borders = OxmlElement("w:pgBorders") # Create new borders element

    # Set all borders
    page_borders.set(qn("w:offsetFrom"), "page")
    for border_name in ["top", "right", "bottom", "left"]:
        border_element = OxmlElement(f"w:{border_name}")
        border_element.set(qn("w:val"), border_type) # Single line
        border_element.set(qn("w:sz"),  str(border_size)) # Line size
        border_element.set(qn("w:space"), str(border_space)) # Space from page sides
        border_element.set(qn("w:color"), border_color) # Line color

        page_borders.append(border_element)
    
    section_properties.append(page_borders)

    return section

def set_margin(section: Section, margin: float = 1.27) -> Section:
    """
    Sets margin to section
    """

    section.top_margin = Cm(margin)
    section.bottom_margin = Cm(margin)
    section.left_margin = Cm(margin)
    section.right_margin = Cm(margin)

    return section

def create_table(document: Document, 
                 rows: int,
                 columns: int) -> Table:
    """
    Creates table in document
    """

    table = document.add_table(rows = rows , cols = columns)
    return table

def set_column_width(table: Table, column_index: int, width: float) -> Table:
    """
    Sets width of column
    """

    for cell in table.columns[column_index].cells:
        cell.width = Cm(width)

    return table

def set_row_height(table: Table, row_index: int, height: int) -> Table:
    """
    Sets height of row in points
    """

    table.rows[row_index].height = Pt(height)

    return table

def merge_cells(cells: List[_Cell]) -> _Cell:
    """
    Merging cells
    """

    merged_cell = cells.pop(0)

    for cell in cells:
        merged_cell = merged_cell.merge(cell)

    return merged_cell

def align_table(table: Table, align: str) -> Table:
    """
    Aligns table
    """

    if (align == "left"):
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
    elif (align == "center"):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    elif (align == "right"):
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    else:
        raise ValueError
    
    return table

def align_paragraph(paragraph: Paragraph, align: str) -> Paragraph:
    """
    Aligns paragraph
    """

    if (align == "left"):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif (align == "center"):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif (align == "right"):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        raise ValueError
    
    return paragraph

def write_text_to_paragraph(paragraph: Paragraph, 
                            text: str,
                            bold: bool = True,
                            underline: bool = False,
                            size: int = 16,
                            name: str = "Cambria") -> Paragraph:
    """
    Writes text to paragraph
    """

    run = paragraph.add_run()

    font = run.font
    font.bold = bold
    font.underline = underline
    font.size = Pt(size)
    font.name = name

    for i, line in enumerate(text.split('\n')):
        run.add_text(line)

        if (i + 1 != len(text.split('\n'))):
            run.add_break()

    return paragraph

def add_qrcode(paragraph: Paragraph, 
               qrcode: BytesIO,
               width: float = 6,
               height: float = 6) -> NoReturn:
    """
    Adds picture to paragraph
    """

    run = paragraph.add_run()
    run.add_picture(qrcode, Cm(width), Cm(height))